import docx
import os
import sqlite3
import sys

from PyQt5 import QtWidgets
from PyQt5.QtGui import QIcon

from window1 import Ui_MainWindow


# Счетчик часов
def s(self, numRows):
    ylist = []
    for row in range(numRows):
        yitem = self.tableWidget.item(row, 1)
        y = int(yitem.text())
        ylist.append(y)
    ysum = str(sum(ylist))
    self.lcd.display(ysum)


class TestApp(Ui_MainWindow):
    def __init__(self, dialog):
        Ui_MainWindow.__init__(self)
        self.setupUi(dialog)

        self.addButton.clicked.connect(self.addRow)  # Кнопка Добавить
        self.editButton.clicked.connect(self.editRow)  # Кнопка Редактировать
        self.deleteButton.clicked.connect(self.deleteRow)  # Кнопка Удалить
        self.saveButton.clicked.connect(self.save)  # Кнопка Схранить табель
        self.updateButton.clicked.connect(self.update)  # Кнопка Загрузить из БД

    def addRow(self):  # Добавляет строку в TableWidget с данными из LineEdit
        x = self.object.text()
        y = self.hour.text()
        if y == '':  # Если не заполнено количество часов
            error_dialog = QtWidgets.QMessageBox()
            error_dialog.setText("Не заполнено количество часов")
            error_dialog.setWindowTitle("ОШИБКА")
            error_dialog.exec_()
        else:
            nummRows = self.tableWidget.rowCount()
            self.tableWidget.insertRow(nummRows)  # Создаем пустую строку
            self.tableWidget.setItem(nummRows, 0, QtWidgets.QTableWidgetItem(x))  # Добавляем текст в строку
            self.tableWidget.setItem(nummRows, 1, QtWidgets.QTableWidgetItem(y))  # Добавляем текст в строку
            self.tableWidget.resizeColumnsToContents()  # Ресайз столбцов под размер введенных данных
            self.object.setText('')  # Стираем данные в LineEdit
            self.hour.setText('')  # Стираем данные в LineEdit
            numRows = self.tableWidget.rowCount()
            s(self, numRows)

    def editRow(self):  # Удаляет строку из TableWidget и добавляет данные в LineEdit
        row = self.tableWidget.currentRow()  # Выделенная строка
        obj = self.tableWidget.item(row, 0).text()
        hou = self.tableWidget.item(row, 1).text()
        self.object.setText(obj)  # Перенос данных из TableWidget в LineEdit
        self.hour.setText(hou)  # Перенос данных из TableWidget в LineEdit
        self.tableWidget.removeRow(row)  # Удаление строки
        numRows = self.tableWidget.rowCount()
        s(self, numRows)

    def deleteRow(self):
        row = self.tableWidget.currentRow()  # Выделенная строка
        self.tableWidget.removeRow(row)  # Удаление строки
        numRows = self.tableWidget.rowCount()
        s(self, numRows)

    def update(self):
        cur_mounth = self.mounth.currentText()  # Значение из ячейки Месяц
        cur_year = self.year.currentText()  # Значение из ячейки Год
        cur_name = self.name.currentText()  # Значение из ячейки Имя
        conn = sqlite3.connect('databases/' + cur_mounth + cur_year + '.db')  # Подключение к базе данных МесяцГод.db
        c = conn.cursor()
        query = 'SELECT * FROM ' + cur_name  # Загрузка данных из таблицы Имя
        result = c.execute(query)
        self.tableWidget.setRowCount(0)  # Перемещение данных в TableWidget
        for row_number, row_data in enumerate(result):
            self.tableWidget.insertRow(row_number)
            for column_number, data in enumerate(row_data):
                self.tableWidget.setItem(row_number, column_number, QtWidgets.QTableWidgetItem(str(data)))
                self.tableWidget.resizeColumnsToContents()
        conn.commit()  # Сохраняем изменения
        c.close()  # Закрываем курсор
        conn.close()  # Закрываем соединение с БД
        numRows = self.tableWidget.rowCount()
        s(self, numRows)

    def save(self):
        cur_mounth = self.mounth.currentText()  # Значение из ячейки Месяц
        cur_year = self.year.currentText()  # Значение из ячейки Год
        cur_name = self.name.currentText()  # Значение из ячейки Имя
        file_path = 'files/' + cur_mounth + cur_year + '.docx'  # Имя создаваемого файла
        if os.path.exists(file_path):  # Если такой файл не существует
            mydoc = docx.Document('files/' + cur_mounth + cur_year + '.docx')
        else:  # Иначе добавляем параграф
            mydoc = docx.Document()
        para = mydoc.add_paragraph(cur_name + ''':
        ''')
        conn = sqlite3.connect('databases/' + cur_mounth + cur_year + '.db')  # Подключение к базе данных МесяцГод.db
        c = conn.cursor()
        c.execute("DROP TABLE IF EXISTS " + cur_name)  # Удаляем таблицу с именем Имя если такая существует
        c.execute("create table if not exists " + cur_name + " (object text, hours text)")    # Создаем таблицу с именем Имя
        for i in range(0, self.tableWidget.rowCount()):
            object = self.tableWidget.item(i, 0).text()
            hours = self.tableWidget.item(i, 1).text()
            c.execute("INSERT into " + cur_name + " VALUES(?, ?)", (object, hours))  # Записываем данные в БД
            para.add_run(object.upper() + ' ' + '(' + hours + ' н/ч), ')  # Записываем данные в файл .docx
        conn.commit()  # Сохраняем изменения
        c.close()  # Закрываем курсор
        conn.close()  # Закрываем соединение с БД
        mydoc.save('files/' + cur_mounth + cur_year + '.docx')  # Сохраняем файл .docx
        self.tableWidget.setRowCount(0)  # Очищаем TableWidget
        self.lcd.display(0)  # Сбрасываем счетчик


if __name__ == '__main__':
    app = QtWidgets.QApplication(sys.argv)
    dialog = QtWidgets.QMainWindow()
    test_1 = TestApp(dialog)
    dialog.show()
    dialog.setWindowTitle('Report v.1.0')
    dialog.setWindowIcon(QIcon('icon.ico'))
    sys.exit(app.exec_())
